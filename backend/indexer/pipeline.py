"""
Document processing pipeline with OCR-first approach.
"""

import os
import logging
from typing import Dict, Any, List, Tuple
from pathlib import Path

try:
    import fitz  # PyMuPDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openai

    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

from .hashutil import get_file_info
from .cache import CacheManager
from .ocr import ocr_processor
from .tables import table_extractor
from .chunk import chunk_processor
from .store import document_store
from . import CONFIG

logger = logging.getLogger(__name__)


class DocumentPipeline:
    def __init__(self):
        self.cache = CacheManager()

        if not OPENAI_AVAILABLE:
            logger.error("OpenAI library not available - embeddings will fail")

    def get_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Get embeddings for text chunks."""
        if not OPENAI_AVAILABLE or not texts:
            return []

        try:
            client = openai.OpenAI()
            response = client.embeddings.create(
                model=CONFIG["EMBED_MODEL"], input=texts
            )

            embeddings = [item.embedding for item in response.data]
            logger.debug(f"Generated {len(embeddings)} embeddings")
            return embeddings

        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")
            return []

    def process_pdf_page(
        self, doc_id: str, pdf_path: str, page_num: int
    ) -> Tuple[str, List[Dict]]:
        """Process a single PDF page with OCR-first approach."""

        # Check cache first
        cached_ocr = self.cache.load_page_ocr(doc_id, page_num)
        cached_tables = self.cache.load_page_tables(doc_id, page_num)

        if cached_ocr is not None and cached_tables is not None:
            logger.debug(f"Using cached data for {doc_id} page {page_num}")
            return cached_ocr, cached_tables

        # Extract text and tables
        page_text = ""
        tables = []

        # Try PyMuPDF text extraction first
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(pdf_path)
                page = doc[page_num - 1]  # PyMuPDF uses 0-based indexing
                page_text = page.get_text()
                doc.close()

                word_count = len(page_text.split()) if page_text else 0
                logger.debug(
                    f"PyMuPDF extracted {word_count} words from page {page_num}"
                )

            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed for page {page_num}: {e}")
                page_text = ""

        # If low word count, use OCR
        word_count = len(page_text.split()) if page_text else 0
        if word_count < CONFIG["OCR_WORD_THRESHOLD"]:
            logger.info(f"Page {page_num} has {word_count} words, using OCR")
            ocr_text, ocr_word_count = ocr_processor.extract_page_text(
                pdf_path, page_num
            )

            if ocr_word_count > word_count:
                page_text = ocr_text
                logger.info(f"OCR improved from {word_count} to {ocr_word_count} words")

        # Extract tables
        if table_extractor.available:
            tables = table_extractor.extract_tables_from_page(pdf_path, page_num)

        # Cache results
        self.cache.save_page_ocr(doc_id, page_num, page_text)
        self.cache.save_page_tables(doc_id, page_num, tables)

        return page_text, tables

    def process_pdf(
        self, file_path: str, doc_id: str, metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Process PDF with OCR-first pipeline."""
        chunks = []

        if not PYMUPDF_AVAILABLE:
            logger.error("PyMuPDF not available - cannot process PDF")
            return chunks

        try:
            doc = fitz.open(file_path)
            total_pages = doc.page_count
            doc.close()

            logger.info(f"Processing PDF {doc_id} with {total_pages} pages")

            for page_num in range(1, total_pages + 1):
                page_text, tables = self.process_pdf_page(doc_id, file_path, page_num)

                # Combine text and table content
                combined_text = page_text
                if tables:
                    table_text = table_extractor.tables_to_text(tables)
                    combined_text = f"{page_text}\n\n{table_text}".strip()

                if combined_text:
                    # Create page metadata
                    page_metadata = {
                        **metadata,
                        "page": page_num,
                        "total_pages": total_pages,
                        "source_type": "pdf_page",
                        "has_tables": len(tables) > 0,
                        "table_count": len(tables),
                    }

                    # Chunk the page content
                    page_chunks = chunk_processor.chunk_text(
                        combined_text, page_metadata
                    )
                    chunks.extend(page_chunks)

                    logger.debug(
                        f"Page {page_num}: {len(page_chunks)} chunks, {len(tables)} tables"
                    )

        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")

        logger.info(f"PDF {doc_id}: {len(chunks)} total chunks")
        return chunks

    def process_docx(
        self, file_path: str, doc_id: str, metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Process DOCX file."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return []

        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)

            if text.strip():
                docx_metadata = {
                    **metadata,
                    "source_type": "docx",
                    "paragraph_count": len(paragraphs),
                }

                return chunk_processor.chunk_text(text, docx_metadata)

        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")

        return []

    def process_text_file(
        self, file_path: str, doc_id: str, metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Process plain text or markdown file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            if text.strip():
                ext = Path(file_path).suffix.lower()
                text_metadata = {
                    **metadata,
                    "source_type": f"text_{ext[1:]}" if ext else "text",
                }

                return chunk_processor.chunk_text(text, text_metadata)

        except Exception as e:
            logger.error(f"Failed to process text file {file_path}: {e}")

        return []

    def process_document(self, file_path: str, watch_dir: str) -> bool:
        """Process a document and index it."""
        try:
            # Get file info
            relative_path, content_hash, doc_id, mtime = get_file_info(
                file_path, watch_dir
            )

            # Check if already cached
            if self.cache.has_cached_data(doc_id, content_hash):
                logger.debug(f"Document {doc_id} already processed")
                return True

            # Base metadata
            base_metadata = {
                "doc_id": doc_id,
                "relative_path": relative_path,
                "content_hash": content_hash,
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "mtime": mtime,
            }

            # Process based on file type
            ext = Path(file_path).suffix.lower()
            chunks = []

            if ext == ".pdf":
                chunks = self.process_pdf(file_path, doc_id, base_metadata)
            elif ext == ".docx":
                chunks = self.process_docx(file_path, doc_id, base_metadata)
            elif ext in {".txt", ".md"}:
                chunks = self.process_text_file(file_path, doc_id, base_metadata)
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return False

            if not chunks:
                logger.warning(f"No chunks extracted from {file_path}")
                return False

            # Generate embeddings
            texts = [chunk["text"] for chunk in chunks]
            embeddings = self.get_embeddings(texts)

            if not embeddings:
                logger.error(f"Failed to generate embeddings for {file_path}")
                return False

            # Delete existing document chunks
            document_store.delete_document(doc_id)

            # Upsert new chunks
            document_store.upsert_chunks(chunks, embeddings)

            # Save document metadata
            doc_meta = {
                **base_metadata,
                "chunk_count": len(chunks),
                "processing_complete": True,
            }
            self.cache.save_doc_meta(doc_id, doc_meta)

            logger.info(f"Successfully indexed {relative_path} ({len(chunks)} chunks)")
            return True

        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {e}")
            return False

    def delete_document(self, file_path: str, watch_dir: str) -> bool:
        """Delete document from index."""
        try:
            relative_path, content_hash, doc_id, _ = get_file_info(file_path, watch_dir)

            # Delete from vector store
            deleted_count = document_store.delete_document(doc_id)

            # Delete cache
            self.cache.delete_doc_cache(doc_id)

            if deleted_count > 0:
                logger.info(
                    f"Deleted document {relative_path} ({deleted_count} chunks)"
                )
            else:
                logger.debug(f"Document {relative_path} not found in index")

            return True

        except Exception as e:
            logger.error(f"Failed to delete document {file_path}: {e}")
            return False


# Global pipeline instance
document_pipeline = DocumentPipeline()
