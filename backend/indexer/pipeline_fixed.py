"""Fixed/Enhanced indexing pipeline."""
import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from .cache import CacheManager
from .chunk_fixed import chunk_processor  # uses the shim we just added
from .config import CONFIG
from .hashutil import get_file_info
from .store import document_store

logger = logging.getLogger(__name__)

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

 # (imports moved to top for E402 compliance)

# Conditional feature imports
if not CONFIG.DISABLE_OCR:
    try:
        from .ocr import ocr_processor
    except ImportError:
        CONFIG.DISABLE_OCR = True
        logger.warning("OCR modules not available, disabling OCR")

if not CONFIG.DISABLE_TABLES:
    try:
        from .tables import table_extractor
    except ImportError:
        CONFIG.DISABLE_TABLES = True
        logger.warning("Table extraction modules not available, disabling tables")

class EnhancedDocumentPipeline:
    """Enhanced document pipeline with full environment flag support."""
    
    def __init__(self):
        self.cache = CacheManager()
        
        # Validate configuration
        errors = CONFIG.validate()
        if errors:
            for error in errors:
                logger.error(f"Configuration error: {error}")
            raise ValueError(f"Configuration validation failed: {errors}")
            
        # Log current configuration
        logger.info(CONFIG.log_settings())
        
        if not OPENAI_AVAILABLE:
            logger.error("OpenAI library not available - embeddings will fail")
            
    def get_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Get embeddings for text chunks."""
        if not OPENAI_AVAILABLE or not texts:
            return []
            
        try:
            client = openai.OpenAI()
            
            # Batch process for efficiency (OpenAI allows up to 2048 inputs)
            batch_size = 100
            all_embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                response = client.embeddings.create(
                    model=CONFIG.EMBED_MODEL,
                    input=batch
                )
                batch_embeddings = [item.embedding for item in response.data]
                all_embeddings.extend(batch_embeddings)
                
            logger.debug(f"Generated {len(all_embeddings)} embeddings in {(len(texts) + batch_size - 1) // batch_size} batches")
            return all_embeddings
            
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")
            return []
    
    def _is_image_only_page(self, text: str, word_count: int) -> bool:
        """Detect if a page contains primarily images with minimal text."""
        if not CONFIG.SKIP_IMAGE_ONLY:
            return False
            
        if word_count < 5:
            return True
            
        # Check character to word ratio (OCR artifacts often have very short "words")
        if word_count > 0:
            avg_word_length = len(text.replace(' ', '')) / word_count
            if avg_word_length < 2:  # Very short average word length suggests OCR artifacts
                return True
                
        # Check for common image-only indicators
        text_lower = text.lower().strip()
        image_only_patterns = [
            r'^\s*\d+\s*$',  # Only page numbers
            r'^\s*[ivxlc]+\s*$',  # Only roman numerals
            r'^\s*figure\s+\d+\s*$',  # Only "Figure N"
        ]
        
        import re
        for pattern in image_only_patterns:
            if re.match(pattern, text_lower):
                return True
                
        return False
        
    def _is_noise_content(self, text: str) -> bool:
        """Filter out obvious noise content."""
        text_lower = text.lower().strip()
        
        if len(text_lower) < 20:  # Too short to be useful
            return True
            
        # Common noise patterns
        noise_patterns = [
            r'intentionally left blank',
            r'this page.*blank',
            r'^\s*\d+\s*$',  # Page numbers only
            r'^\s*copyright.*all rights reserved\s*$',
            r'^\s*confidential\s*$',
            r'^\s*draft\s*$',
        ]
        
        import re
        for pattern in noise_patterns:
            if re.search(pattern, text_lower):
                return True
                
        return False
            
    def process_pdf_page(self, doc_id: str, pdf_path: str, page_num: int) -> Tuple[str, List[Dict]]:
        """Process a single PDF page with environment flag support."""
        
        # Check cache first
        cached_ocr = self.cache.load_page_ocr(doc_id, page_num)
        cached_tables = self.cache.load_page_tables(doc_id, page_num)
        
        if cached_ocr is not None and cached_tables is not None:
            logger.debug(f"Using cached data for {doc_id} page {page_num}")
            return cached_ocr, cached_tables
            
        # Extract text and tables
        page_text = ""
        tables = []
        
        # Try PyMuPDF text extraction first
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(pdf_path)
                page = doc[page_num - 1]  # PyMuPDF uses 0-based indexing
                page_text = page.get_text()
                doc.close()
                
                word_count = len(page_text.split()) if page_text else 0
                logger.debug(f"PyMuPDF extracted {word_count} words from page {page_num}")
                
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed for page {page_num}: {e}")
                page_text = ""
                
        # OCR fallback (if enabled and needed)
        word_count = len(page_text.split()) if page_text else 0
        if not CONFIG.DISABLE_OCR and word_count < CONFIG.OCR_WORD_THRESHOLD:
            logger.info(f"Page {page_num} has {word_count} words, using OCR (OCR enabled: {not CONFIG.DISABLE_OCR})")
            try:
                ocr_text, ocr_word_count = ocr_processor.extract_page_text(pdf_path, page_num)
                
                if ocr_word_count > word_count:
                    page_text = ocr_text
                    word_count = ocr_word_count
                    logger.info(f"OCR improved from {word_count} to {ocr_word_count} words")
            except Exception as e:
                logger.error(f"OCR failed for page {page_num}: {e}")
                
        # Check if page should be skipped (image-only detection)
        if self._is_image_only_page(page_text, word_count):
            logger.info(f"Skipping image-only page {page_num}")
            return "", []
            
        # Extract tables (if enabled)
        if not CONFIG.DISABLE_TABLES and 'table_extractor' in globals():
            try:
                tables = table_extractor.extract_tables_from_page(pdf_path, page_num)
                logger.debug(f"Extracted {len(tables)} tables from page {page_num}")
            except Exception as e:
                logger.error(f"Table extraction failed for page {page_num}: {e}")
                
        # Cache results
        self.cache.save_page_ocr(doc_id, page_num, page_text)
        self.cache.save_page_tables(doc_id, page_num, tables)
        
        return page_text, tables
        
    def process_pdf(self, file_path: str, doc_id: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Process PDF with enhanced pipeline and flag support."""
        chunks = []
        
        if not PYMUPDF_AVAILABLE:
            logger.error("PyMuPDF not available - cannot process PDF")
            return chunks
            
        try:
            doc = fitz.open(file_path)
            total_pages = doc.page_count
            doc.close()
            
            logger.info(f"Processing PDF {doc_id} with {total_pages} pages (OCR: {'disabled' if CONFIG.DISABLE_OCR else 'enabled'}, Tables: {'disabled' if CONFIG.DISABLE_TABLES else 'enabled'})")
            
            # Process all pages and collect text for header/footer detection
            page_texts = []
            page_tables = []
            
            for page_num in range(1, total_pages + 1):
                page_text, tables = self.process_pdf_page(doc_id, file_path, page_num)
                page_texts.append(page_text)
                page_tables.append(tables)
                
            # Detect and remove common headers/footers
            headers_footers = self._detect_headers_footers(page_texts)
            
            # Extract document structure for TOC/heading-aware indexing
            pages_text_dict = {}
            cleaned_texts = []
            for page_num, page_text in enumerate(page_texts, 1):
                cleaned_text = self._clean_headers_footers(page_text, headers_footers)
                cleaned_texts.append(cleaned_text)
                if cleaned_text.strip():  # Only include non-empty pages
                    pages_text_dict[page_num] = cleaned_text
            
            # Build heading index from document structure
            from .structure import extract_structure_for_document
            heading_index = extract_structure_for_document(file_path, pages_text_dict)
            
            # Create chunks from cleaned pages
            for page_num, (cleaned_text, tables) in enumerate(zip(cleaned_texts, page_tables), 1):
                if not cleaned_text and not tables:
                    continue
                
                # Filter noise
                if self._is_noise_content(cleaned_text):
                    logger.debug(f"Skipping noise content on page {page_num}")
                    continue
                
                # Combine text and table content
                combined_text = cleaned_text
                if tables and not CONFIG.DISABLE_TABLES:
                    table_text = self._tables_to_text(tables)
                    combined_text = f"{cleaned_text}\n\n{table_text}".strip()
                    
                if combined_text.strip():
                    # Create page metadata
                    page_metadata = {
                        **metadata,
                        "page": page_num,
                        "total_pages": total_pages,
                        "source_type": "pdf_page",
                        "has_tables": len(tables) > 0,
                        "table_count": len(tables),
                        "ocr_used": not CONFIG.DISABLE_OCR,
                        "tables_extracted": not CONFIG.DISABLE_TABLES
                    }
                    
                    # Chunk the page content and enrich with structural metadata
                    page_chunks = chunk_processor.chunk_text(combined_text, page_metadata)
                    
                    # Enrich chunks with structural metadata
                    from .structure import tag_chunk
                    for i, chunk in enumerate(page_chunks):
                        structural_meta = tag_chunk(chunk["text"], page_num, i, heading_index)
                        if structural_meta:
                            chunk["metadata"].update(structural_meta)
                    
                    chunks.extend(page_chunks)
                    
                    logger.debug(f"Page {page_num}: {len(page_chunks)} chunks, {len(tables)} tables")
                    
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
            
        logger.info(f"PDF {doc_id}: {len(chunks)} total chunks")
        return chunks
    
    def _detect_headers_footers(self, page_texts: List[str]) -> Dict[str, str]:
        """Detect common headers and footers across pages."""
        if len(page_texts) < 3:  # Need at least 3 pages to detect patterns
            return {"header": "", "footer": ""}
            
        # Extract first and last lines from each page
        first_lines = []
        last_lines = []
        
        for text in page_texts:
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            if lines:
                first_lines.append(lines[0][:100])  # First 100 chars
                last_lines.append(lines[-1][:100])  # Last 100 chars
                
        # Find most common patterns
        from collections import Counter
        
        first_counter = Counter(first_lines)
        last_counter = Counter(last_lines)
        
        # A header/footer should appear on at least 30% of pages
        min_occurrences = max(2, len(page_texts) * 0.3)
        
        header = ""
        footer = ""
        
        if first_counter:
            most_common_first, count = first_counter.most_common(1)[0]
            if count >= min_occurrences:
                header = most_common_first
                
        if last_counter:
            most_common_last, count = last_counter.most_common(1)[0]
            if count >= min_occurrences:
                footer = most_common_last
                
        if header or footer:
            logger.debug(f"Detected header: '{header[:50]}...', footer: '{footer[:50]}...'")
            
        return {"header": header, "footer": footer}
    
    def _clean_headers_footers(self, text: str, headers_footers: Dict[str, str]) -> str:
        """Remove detected headers and footers from text."""
        header = headers_footers.get("header", "")
        footer = headers_footers.get("footer", "")
        
        if header and text.startswith(header):
            text = text[len(header):].lstrip('\n ')
            
        if footer and text.endswith(footer):
            text = text[:-len(footer)].rstrip('\n ')
            
        return text
    
    def _tables_to_text(self, tables: List[Dict]) -> str:
        """Convert table data to text format."""
        if not tables:
            return ""
            
        table_texts = []
        for i, table in enumerate(tables):
            table_text = f"Table {i+1}:\n"
            # Add table content (implementation depends on table format)
            table_text += str(table.get('data', ''))  # Simplified
            table_texts.append(table_text)
            
        return "\n\n".join(table_texts)
        
    def process_spreadsheet(self, file_path: str, doc_id: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Process Excel/CSV files by flattening to text."""
        if not PANDAS_AVAILABLE:
            logger.error("pandas not available - cannot process spreadsheets")
            return []
            
        try:
            ext = Path(file_path).suffix.lower()
            
            if ext == '.csv':
                df = pd.read_csv(file_path)
                sheets_data = {'Sheet1': df}
            else:  # Excel files
                sheets_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
                
            all_text_parts = []
            for sheet_name, df in sheets_data.items():
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False, max_rows=None)
                all_text_parts.append(sheet_text)
                
            combined_text = "\n\n".join(all_text_parts)
            
            if combined_text.strip():
                spreadsheet_metadata = {
                    **metadata,
                    "source_type": f"spreadsheet_{ext[1:]}",
                    "sheet_count": len(sheets_data),
                    "sheet_names": list(sheets_data.keys())
                }
                
                return chunk_processor.chunk_text(combined_text, spreadsheet_metadata)
                
        except Exception as e:
            logger.error(f"Failed to process spreadsheet {file_path}: {e}")
            
        return []
        
    def process_docx(self, file_path: str, doc_id: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Process DOCX file."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return []
            
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            
            if text.strip():
                docx_metadata = {
                    **metadata,
                    "source_type": "docx",
                    "paragraph_count": len(paragraphs)
                }
                
                return chunk_processor.chunk_text(text, docx_metadata)
                
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
            
        return []
        
    def process_text_file(self, file_path: str, doc_id: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Process plain text or markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                
            if text.strip():
                ext = Path(file_path).suffix.lower()
                text_metadata = {
                    **metadata,
                    "source_type": f"text_{ext[1:]}" if ext else "text",
                }
                
                return chunk_processor.chunk_text(text, text_metadata)
                
        except Exception as e:
            logger.error(f"Failed to process text file {file_path}: {e}")
            
        return []
        
    def process_document(self, file_path: str, watch_dir: str) -> bool:
        """Process a document and index it with full environment flag support."""
        try:
            # Get file info
            relative_path, content_hash, doc_id, mtime = get_file_info(file_path, watch_dir)
            
            # Check file extension against allowed list
            ext = Path(file_path).suffix.lower()
            if ext not in CONFIG.ALLOWED_EXT:
                logger.debug(f"Skipping {file_path}: extension {ext} not in allowed list")
                return False
            
            # Check if already cached
            if self.cache.has_cached_data(doc_id, content_hash):
                logger.debug(f"Document {doc_id} already processed")
                return True
                
            # Base metadata
            base_metadata = {
                "doc_id": doc_id,
                "relative_path": relative_path,
                "content_hash": content_hash,
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "mtime": mtime,
                "processing_config": {
                    "ocr_disabled": CONFIG.DISABLE_OCR,
                    "tables_disabled": CONFIG.DISABLE_TABLES,
                    "skip_image_only": CONFIG.SKIP_IMAGE_ONLY,
                    "chunk_size": CONFIG.CHUNK_CHARS,
                    "chunk_overlap": CONFIG.CHUNK_OVERLAP
                }
            }
            
            # Process based on file type
            chunks = []
            
            if ext == ".pdf":
                chunks = self.process_pdf(file_path, doc_id, base_metadata)
            elif ext in {".xlsx", ".xls", ".csv"}:
                chunks = self.process_spreadsheet(file_path, doc_id, base_metadata)
            elif ext in {".docx", ".doc"}:
                chunks = self.process_docx(file_path, doc_id, base_metadata)
            elif ext in {".txt", ".md", ".rtf"}:
                chunks = self.process_text_file(file_path, doc_id, base_metadata)
            elif ext in {".pptx", ".ppt"}:
                logger.warning(f"PowerPoint processing not yet implemented: {ext}")
                return False
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return False
                
            if not chunks:
                logger.warning(f"No chunks extracted from {file_path}")
                return False
                
            # Generate embeddings
            texts = [chunk["text"] for chunk in chunks]
            embeddings = self.get_embeddings(texts)
            
            if not embeddings:
                logger.error(f"Failed to generate embeddings for {file_path}")
                return False
                
            # Delete existing document chunks
            document_store.delete_document(doc_id)
            
            # Upsert new chunks
            document_store.upsert_chunks(chunks, embeddings)
            
            # Save document metadata
            doc_meta = {
                **base_metadata,
                "chunk_count": len(chunks),
                "processing_complete": True
            }
            self.cache.save_doc_meta(doc_id, doc_meta)
            
            logger.info(f"Successfully indexed {relative_path} ({len(chunks)} chunks)")
            return True
            
        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {e}")
            return False
            
    def delete_document(self, file_path: str, watch_dir: str) -> bool:
        """Delete document from index."""
        try:
            relative_path, content_hash, doc_id, _ = get_file_info(file_path, watch_dir)
            
            # Delete from vector store
            deleted_count = document_store.delete_document(doc_id)
            
            # Delete cache
            self.cache.delete_doc_cache(doc_id)
            
            if deleted_count > 0:
                logger.info(f"Deleted document {relative_path} ({deleted_count} chunks)")
            else:
                logger.debug(f"Document {relative_path} not found in index")
                
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document {file_path}: {e}")
            return False


# Global pipeline instance
document_pipeline = EnhancedDocumentPipeline()
